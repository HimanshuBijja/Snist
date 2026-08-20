#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Extract readable text from every source file in a folder.

Usage:
    python extract_sources.py <source_folder> <output_folder>

Handles .pdf (via pdftotext), .docx (via python-docx), .pptx (via zip+XML,
since python-pptx is not installed), and passes through .txt/.md.
Skips the blueprint image and instructions.md.

Writes one <name>.txt per source file, plus prints a summary so you can see
line counts and confirm nothing was silently missed.
"""
import os
import re
import subprocess
import sys
import zipfile

SKIP_NAMES = {'instructions.md', 'answers.md'}
IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}


def from_pdf(path, out):
    subprocess.run(['pdftotext', '-layout', path, out], check=True)
    return open(out, encoding='utf-8', errors='replace').read()


def from_docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:                     # tables hold real content too
        for row in table.rows:
            parts.append(' | '.join(c.text.strip() for c in row.cells))
    return '\n'.join(parts)


def from_pptx(path):
    """python-pptx is unavailable, so read the OOXML directly.
    A .pptx is a zip; slide text lives in <a:t> elements of ppt/slides/slideN.xml.
    Speaker notes live in ppt/notesSlides/notesSlideN.xml."""
    out = []
    with zipfile.ZipFile(path) as z:
        def num(n):
            m = re.search(r'(\d+)\.xml$', n)
            return int(m.group(1)) if m else 0

        slides = sorted((n for n in z.namelist()
                         if re.match(r'ppt/slides/slide\d+\.xml$', n)), key=num)
        notes = {num(n): n for n in z.namelist()
                 if re.match(r'ppt/notesSlides/notesSlide\d+\.xml$', n)}

        for name in slides:
            i = num(name)
            xml = z.read(name).decode('utf-8', 'replace')
            texts = re.findall(r'<a:t>(.*?)</a:t>', xml, re.S)
            out.append('--- Slide %d ---' % i)
            out.extend(unescape(t) for t in texts if t.strip())
            if i in notes:
                nxml = z.read(notes[i]).decode('utf-8', 'replace')
                ntexts = [unescape(t) for t in
                          re.findall(r'<a:t>(.*?)</a:t>', nxml, re.S) if t.strip()]
                # the slide-number placeholder shows up as a lone digit; drop it
                ntexts = [t for t in ntexts if not t.strip().isdigit()]
                if ntexts:
                    out.append('[Notes] ' + ' '.join(ntexts))
            out.append('')
    return '\n'.join(out)


def unescape(s):
    return (s.replace('&amp;', '&').replace('&lt;', '<')
             .replace('&gt;', '>').replace('&quot;', '"').replace('&#39;', "'"))


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    src, dst = sys.argv[1], sys.argv[2]
    os.makedirs(dst, exist_ok=True)

    results, skipped = [], []
    for fn in sorted(os.listdir(src)):
        path = os.path.join(src, fn)
        if not os.path.isfile(path):
            continue
        stem, ext = os.path.splitext(fn)
        ext = ext.lower()

        if fn.lower() in SKIP_NAMES:
            skipped.append((fn, 'instructions/output file'))
            continue
        if ext in IMAGE_EXTS:
            note = 'BLUEPRINT - read this with the Read tool' \
                if stem.lower().startswith('blueprint') else 'image'
            skipped.append((fn, note))
            continue

        out = os.path.join(dst, stem + '.txt')
        try:
            if ext == '.pdf':
                text = from_pdf(path, out)
            elif ext == '.docx':
                text = from_docx(path)
            elif ext == '.pptx':
                text = from_pptx(path)
            elif ext in ('.txt', '.md'):
                text = open(path, encoding='utf-8', errors='replace').read()
            else:
                skipped.append((fn, 'unsupported extension'))
                continue
        except Exception as e:
            skipped.append((fn, 'FAILED: %s: %s' % (type(e).__name__, e)))
            continue

        with open(out, 'w', encoding='utf-8') as f:
            f.write(text)
        results.append((fn, stem + '.txt', text.count('\n') + 1, len(text)))

    print('EXTRACTED %d file(s) -> %s\n' % (len(results), dst))
    for fn, out, lines, chars in results:
        print('  %-34s -> %-28s %5d lines  %7d chars' % (fn, out, lines, chars))
    if skipped:
        print('\nSKIPPED:')
        for fn, why in skipped:
            print('  %-34s    %s' % (fn, why))
    print('\nRead EVERY .txt above in full before writing answers.md.')


if __name__ == '__main__':
    main()
